import os
from datetime import datetime, timezone

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.modules.report_generator.schemas import ReportResponse


def build_docx_report(report: ReportResponse, mode: str, output_path: str) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h1_style = doc.styles["Heading 1"]
    h1_style.font.color.rgb = None

    title = report.title or "Creative Pre-Test Report"
    doc.add_heading(title, level=0)

    mode_label = "Internal Report" if mode == "internal" else "Client-Facing Report"
    p = doc.add_paragraph()
    p.add_run(mode_label).italic = True

    _add_section(doc, "Executive Summary", report.summary or "No summary available.")

    _add_section(doc, "Creative Overview",
        f"Test Run ID: {report.test_run_id}\n"
        f"Status: Completed\n"
        f"Generated: {datetime.now(timezone.utc).isoformat()}")

    _add_section(doc, "Main Message", report.main_message or "No main message extracted.")

    if report.audience_reactions:
        doc.add_heading("Audience Reactions", level=1)
        for a in report.audience_reactions:
            doc.add_heading(a.segment_name or a.audience_profile_id, level=2)
            doc.add_paragraph(f"Reaction: {a.reaction}")
            doc.add_paragraph(f"Engagement probability: {a.engagement_probability:.0%}")
            if a.positive_triggers:
                doc.add_paragraph(f"Positive triggers: {', '.join(a.positive_triggers)}")
            if a.objections:
                doc.add_paragraph(f"Objections: {', '.join(a.objections)}")

    if report.scorecard:
        doc.add_heading("Scorecard", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Criterion"
        hdr[1].text = "Score"
        for s in report.scorecard:
            row = table.add_row().cells
            row[0].text = s.criterion
            row[1].text = f"{s.score:.0f}/10"
        avg = sum(s.score for s in report.scorecard) / max(len(report.scorecard), 1)
        doc.add_paragraph(f"Average score: {avg:.1f}/10")

    if report.risks:
        doc.add_heading("Brand Safety and Risks", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Severity"
        hdr[1].text = "Risk"
        for r in report.risks:
            row = table.add_row().cells
            row[0].text = r.severity.upper()
            row[1].text = r.risk
    else:
        _add_section(doc, "Brand Safety and Risks", "No significant risks identified.")

    _add_section(doc, "Visual Analysis Notes", report.visual_notes or "No visual notes.")

    if report.recommendations:
        doc.add_heading("Recommendations", level=1)
        for r in report.recommendations:
            p = doc.add_paragraph()
            p.add_run(f"[{r.priority.upper()}] ").bold = True
            p.add_run(r.recommendation)

    _add_section(doc, "Final Recommendation", report.final_recommendation.upper() if report.final_recommendation else "REVISE")

    if mode == "internal":
        _add_section(doc, "Appendix: Test Context",
            f"Report ID: {report.id}\n"
            f"Mode: {mode}\n"
            f"Format: {report.report_format}\n"
            f"Version: {report.version}")

    doc.save(output_path)
    return output_path


def _add_section(doc: Document, title: str, content: str):
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
