"""Tests for DOCX export generation."""

import os

from httpx import ASGITransport, AsyncClient
from src.main import app

from src.modules.export_jobs.renderers.docx_report import build_docx_report
from src.modules.report_generator.schemas import ReportResponse
from src.modules.report_generator.service import generate_report
from src.modules.creative_assets.service import create_asset as create_creative_asset
from src.modules.creative_assets.schemas import CreateCreativeAssetRequest
from src.modules.test_runs.schemas import CreateTestRunRequest
from src.modules.test_runs.service import create_test_run, run_test_run
from src.shared.config.settings import get_settings


def test_docx_renderer_creates_file():
    asset = create_creative_asset(CreateCreativeAssetRequest(
        title="DOCX Test Asset", asset_type="text", text_content="Test content for DOCX."
    ))
    run = create_test_run(CreateTestRunRequest(creative_asset_id=asset.id))
    run = run_test_run(run.id)
    report = generate_report(run.id, report_mode="internal", report_format="json")
    assert report is not None

    settings = get_settings()
    output_path = os.path.join(settings.exports_root, "test_docx.docx")
    os.makedirs(settings.exports_root, exist_ok=True)

    result = build_docx_report(report, "internal", output_path)
    assert os.path.isfile(result)
    assert os.path.getsize(result) > 0
    os.remove(result)


def test_docx_export_endpoint_creates_job():
    transport = ASGITransport(app=app)
    import anyio

    async def test():
        asset = create_creative_asset(CreateCreativeAssetRequest(
            title="DOCX Endpoint Asset", asset_type="text", text_content="DOCX endpoint test."
        ))
        run = create_test_run(CreateTestRunRequest(creative_asset_id=asset.id))
        run = run_test_run(run.id)
        report = generate_report(run.id, report_mode="internal", report_format="json")
        assert report is not None

        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.post(f"/exports/report/{report.id}/docx")
        assert resp.status_code == 201
        data = resp.json()
        assert data["export_type"] == "docx"
        assert data["status"] == "completed"
        assert data["file_path"] is not None

    anyio.run(test)


def test_docx_file_exists_and_nonzero():
    transport = ASGITransport(app=app)
    import anyio

    async def test():
        asset = create_creative_asset(CreateCreativeAssetRequest(
            title="DOCX File Test", asset_type="text", text_content="Test content."
        ))
        run = create_test_run(CreateTestRunRequest(creative_asset_id=asset.id))
        run = run_test_run(run.id)
        report = generate_report(run.id, report_mode="internal", report_format="json")

        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.post(f"/exports/report/{report.id}/docx")
        assert resp.status_code == 201
        data = resp.json()
        assert os.path.isfile(data["file_path"])
        assert os.path.getsize(data["file_path"]) > 100

    anyio.run(test)


def test_docx_report_has_sections():
    asset = create_creative_asset(CreateCreativeAssetRequest(
        title="DOCX Sections Asset", asset_type="text", text_content="Multiple sections test."
    ))
    run = create_test_run(CreateTestRunRequest(creative_asset_id=asset.id))
    run = run_test_run(run.id)
    report = generate_report(run.id, report_mode="internal", report_format="json")
    assert report is not None

    settings = get_settings()
    output_path = os.path.join(settings.exports_root, "test_docx_sections.docx")
    os.makedirs(settings.exports_root, exist_ok=True)
    build_docx_report(report, "internal", output_path)

    from docx import Document
    doc = Document(output_path)
    texts = [p.text for p in doc.paragraphs]
    combined = " ".join(texts)
    assert "Executive Summary" in combined or "Summary" in combined
    assert "Final Recommendation" in combined
    os.remove(output_path)


def test_docx_client_facing_mode():
    asset = create_creative_asset(CreateCreativeAssetRequest(
        title="DOCX Client Asset", asset_type="text", text_content="Client-facing test."
    ))
    run = create_test_run(CreateTestRunRequest(creative_asset_id=asset.id))
    run = run_test_run(run.id)
    report = generate_report(run.id, report_mode="client_facing", report_format="json")
    assert report is not None

    settings = get_settings()
    output_path = os.path.join(settings.exports_root, "test_docx_client.docx")
    os.makedirs(settings.exports_root, exist_ok=True)
    build_docx_report(report, "client_facing", output_path)

    from docx import Document
    doc = Document(output_path)
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "Client" in texts or "Recommendation" in texts
    os.remove(output_path)
